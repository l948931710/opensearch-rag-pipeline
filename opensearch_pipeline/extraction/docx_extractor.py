# -*- coding: utf-8 -*-
"""
docx_extractor.py — Word/DOCX 文本提取器

使用 python-docx 提取段落和表格，生成结构化 blocks。
逻辑 refactored from scan_pending_clean.py:L248-L265。

标题检测策略（用户决策：用 style + regex fallback）：
  1. 优先用 paragraph.style.name 判断 heading（Heading 1 / Heading 2）
  2. 如果 style 是 Normal 但文本匹配中文标题模式 → 也标记为 heading
  3. 工厂文档可能不使用标准 Word 样式，regex fallback 确保兼容

生产依赖：python-docx
模拟模式：不需要真实 DOCX 文件
"""

import re
from typing import TYPE_CHECKING, List, Optional, Tuple

from opensearch_pipeline.extraction.schema import ExtractedBlock, is_pseudo_heading

if TYPE_CHECKING:  # resolve the "ImageAsset" forward-ref annotation without a runtime import
    from opensearch_pipeline.extraction.image_extraction_utils import ImageAsset

# Word 标题样式名 → heading level 映射
_STYLE_LEVEL_MAP = {
    "Heading 1": 1, "Heading 2": 2, "Heading 3": 3, "Heading 4": 4,
    "heading 1": 1, "heading 2": 2, "heading 3": 3, "heading 4": 4,
    "标题 1": 1, "标题 2": 2, "标题 3": 3, "标题 4": 4,
    "Title": 1, "Subtitle": 2,
}

# 中文标题正则 fallback
_CN_HEADING_RE = re.compile(
    r"^(?:"
    r"第[一二三四五六七八九十\d]+[章节条款部分]\s*.+|"
    r"[一二三四五六七八九十]+[、\.]\s*.+|"
    r"（[一二三四五六七八九十\d]+）\s*.+"
    r")$"
)
_SUB_HEADING_RE = re.compile(r"^\d+\.\d+\s+.+$")

# 标题最大字数：超过即视为正文段落（与下方 regex fallback 的同一阈值一致）。
_HEADING_MAX_LEN = 30
# 行尾断句标点：真正的标题极少以句末标点收尾，正文/条款句几乎总是。
_SENTENCE_FINAL = "。；！？．"
# 平铺枚举条款前缀（阿拉伯数字 + 右括号/顿号）："1）" "2)" "3、"。这类是条款/列表
# 正文项，不是章节标题（与 _CN_HEADING_RE 故意只认中文数字一致）。点分层级号
# （"5.1" "4.2.1"）是真子标题 → 由 _DOTTED_SECTION_RE 豁免。
_ENUM_CLAUSE_PREFIX_RE = re.compile(r"^\s*\d+\s*[)）、]")
_DOTTED_SECTION_RE = re.compile(r"^\s*\d+\.\d")


def _looks_like_clause_body(text: str) -> bool:
    """style 把一段文字判成 heading，但文字本身明显是正文/条款句时返回 True，
    用于否决其 heading 身份。

    动机：部分 制度/管理规定 Word 文档把正文段落滥用了 Subtitle / 标题样式
    （it_IT信息系统内控制度.docx 实证：62 块里 48 块被 Subtitle 样式误判 heading，
    下游 clause 切块把 heading 文字当 section_label 而非正文 → 82% 正文 collapse，
    recall 0.30）。仅靠样式信任会丢正文；此处对样式判定再加一道"内容是否像标题"
    的复核。判定为正文的三类形状：
      1) 圈数字标注（与 schema.is_pseudo_heading 同口径）；
      2) 平铺阿拉伯枚举条款项（"1）申请…"），点分层级号 "5.1" 除外；
      3) 过长（> _HEADING_MAX_LEN）或以句末标点收尾的整句。
    """
    if not text:
        return False
    # 1) 圈数字标注 callout
    if is_pseudo_heading(text):
        return True
    # 2) 平铺阿拉伯枚举条款项（点分层级子标题豁免）
    if _ENUM_CLAUSE_PREFIX_RE.match(text) and not _DOTTED_SECTION_RE.match(text):
        return True
    # 3) 过长 → 正文段落；或句末标点收尾 → 是句子而非标签
    if len(text) > _HEADING_MAX_LEN:
        return True
    if text[-1] in _SENTENCE_FINAL:
        return True
    return False


def _detect_heading_level(style_name: str, text: str) -> Optional[int]:
    """
    检测标题级别。

    策略：style 优先，regex fallback。style 命中后仍用 _looks_like_clause_body
    复核内容，避免"正文滥用 Subtitle/标题样式"的制度文档丢正文（见该函数 docstring）。
    """
    stripped = text.strip()

    # 1. Word 样式检测
    style_level: Optional[int] = None
    if style_name in _STYLE_LEVEL_MAP:
        style_level = _STYLE_LEVEL_MAP[style_name]
    # 2. 样式名包含 "Heading" 或 "标题"
    elif "heading" in style_name.lower():
        nums = re.findall(r"\d+", style_name)
        style_level = min(int(nums[0]), 4) if nums else 2

    if style_level is not None:
        # 样式声称是标题，但文字本身是正文/条款句 → 回落为 paragraph，
        # 否则 clause 切块会把这段正文当 section_label 丢弃。
        if _looks_like_clause_body(stripped):
            return None
        return style_level

    # 3. Regex fallback：中文标题模式 (限制最大长度以防长正文段落被误判为标题而遗漏)
    # 标注式 callout（圈数字开头）veto —— 与 pdf_extractor 同口径；现有两条
    # 正则形状本就不可能匹配圈数字开头，guard 防未来放宽正则时回归。
    if is_pseudo_heading(stripped):
        return None
    if len(stripped) <= 30:
        if _CN_HEADING_RE.match(stripped):
            if stripped.startswith("第"):
                return 1
            return 2

        if _SUB_HEADING_RE.match(stripped):
            return 3

    return None


def extract_docx(
    local_path: str,
) -> Tuple[List[ExtractedBlock], List[str]]:
    """
    从 DOCX 文件提取 blocks。

    使用递归的顺序块提取器来保留段落与表格的自然物理阅读顺序，
    并自动解包作为边框装饰的单单元格（1x1）外表层表格。

    Returns:
        (blocks, warnings)
    """
    try:
        import docx
        from docx.text.paragraph import Paragraph
        from docx.table import Table
    except ImportError:
        return [], ["python-docx not installed, cannot extract DOCX"]

    warnings: List[str] = []
    blocks: List[ExtractedBlock] = []
    current_section_ref = [None]  # 用列表传递可变引用，跟踪当前所属 section_path

    try:
        document = docx.Document(local_path)
    except Exception as e:
        return [], [f"Failed to open DOCX: {e}"]

    def _extract_recursive(parent) -> List[ExtractedBlock]:
        extracted = []
        if hasattr(parent, 'element') and hasattr(parent.element, 'body'):
            parent_elm = parent.element.body
        elif hasattr(parent, '_tc'):
            parent_elm = parent._tc
        elif hasattr(parent, 'element'):
            parent_elm = parent.element
        else:
            parent_elm = parent

        table_idx = 0
        for child in parent_elm.iterchildren():
            if child.tag.endswith('p'):
                para = Paragraph(child, parent)
                text = para.text.strip()
                tb_texts = _textbox_texts(child)
                if not text and not tb_texts:
                    continue

                style_name = para.style.name if para.style else "Normal"
                heading_level = _detect_heading_level(style_name, text) if text else None

                if text and heading_level is not None:
                    current_section_ref[0] = text
                    extracted.append(ExtractedBlock(
                        block_type="heading",
                        text=text,
                        level=heading_level,
                        section_path=current_section_ref[0],
                        source="native",
                        extra={"word_style": style_name},
                    ))
                elif text:
                    extracted.append(ExtractedBlock(
                        block_type="paragraph",
                        text=text,
                        section_path=current_section_ref[0],
                        source="native",
                        extra={"word_style": style_name},
                    ))
                # 文本框内容（para.text 不含）→ 独立 paragraph block
                for tb in tb_texts:
                    extracted.append(ExtractedBlock(
                        block_type="paragraph",
                        text=tb,
                        section_path=current_section_ref[0],
                        source="native",
                        extra={"word_style": style_name, "from_textbox": True},
                    ))
            elif child.tag.endswith('tbl'):
                table = Table(child, parent)
                # 识别单单元格 (1x1) 装饰性外表格并递归解包其内部子段落/子表格
                if len(table.rows) == 1 and len(table.rows[0].cells) == 1:
                    cell = table.rows[0].cells[0]
                    extracted.extend(_extract_recursive(cell))
                else:
                    rows_text = []
                    # DC-3: dedup merged cells (gridSpan/vMerge) — python-docx repeats the same
                    # <w:tc> across spanned grid positions/rows. Store the element itself (not id():
                    # lxml proxy ids get GC-reused → distinct cells collide & get dropped); holding a
                    # reference keeps lxml's per-node proxy stable so membership is true identity.
                    seen_tc = set()
                    for row in table.rows:
                        cells = []
                        for cell in row.cells:
                            tc = getattr(cell, "_tc", None)
                            if tc is not None:
                                if tc in seen_tc:
                                    continue
                                seen_tc.add(tc)
                            cell_text = _cell_text_with_textboxes(cell)
                            if cell_text:
                                cells.append(cell_text)
                        if cells:
                            rows_text.append(" | ".join(cells))

                    if rows_text:
                        table_md = "\n".join(f"| {row} |" for row in rows_text)
                        extracted.append(ExtractedBlock(
                            block_type="table",
                            text=table_md,
                            section_path=current_section_ref[0],
                            source="native",
                            extra={"table_index": table_idx, "row_count": len(rows_text)},
                        ))
                    table_idx += 1
        return extracted

    try:
        blocks = _extract_recursive(document)
    except Exception as e:
        warnings.append(f"Recursive extraction encountered an issue: {e}")
        # fallback to basic non-recursive paragraphs extraction to be safe
        blocks = []
        for para in document.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style_name = para.style.name if para.style else "Normal"
            blocks.append(ExtractedBlock(
                block_type="paragraph",
                text=text,
                section_path=None,
                source="native",
                extra={"word_style": style_name},
            ))

    return blocks, warnings


# ═══════════════════════════════════════════════════════════════
# 方案 B — 保留图片位置的增强提取
# ═══════════════════════════════════════════════════════════════

# OOXML 命名空间常量
_NSMAP_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_NSMAP_WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
_NSMAP_V = "urn:schemas-microsoft-com:vml"
_NSMAP_R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

# XPath 用于在 <w:p> 中查找嵌入图片引用
# <w:drawing> 包裹 <wp:inline> 或 <wp:anchor>，内含 <a:blip r:embed="rIdXX">
_XPATH_DRAWING = f".//{{{_NSMAP_W}}}drawing"
_XPATH_PICT = f".//{{{_NSMAP_W}}}pict"

# <a:blip> 里的 r:embed 属性包含 rId
_NSMAP_A = "http://schemas.openxmlformats.org/drawingml/2006/main"
_XPATH_BLIP = f".//{{{_NSMAP_A}}}blip"

# VML <v:imagedata> 的 r:id 属性（旧式 <w:pict> 路径）
_XPATH_IMAGEDATA = f".//{{{_NSMAP_V}}}imagedata"


def _find_image_rel_ids(para_element) -> List[str]:
    """
    从一个 <w:p> XML 元素中提取所有嵌入图片的 relationship ID。

    检测两种嵌入方式：
      1. <w:drawing> → <a:blip r:embed="rIdXX">  （现代 DrawingML 路径）
      2. <w:pict>    → <v:imagedata r:id="rIdXX"> （旧版 VML 路径）

    Returns:
        rel_id 字符串列表，按出现顺序排列（通常一个段落只有一张图）。
    """
    rel_ids: List[str] = []
    r_embed = f"{{{_NSMAP_R}}}embed"
    r_id_attr = f"{{{_NSMAP_R}}}id"

    # 路径 1: DrawingML — <w:drawing> / <a:blip r:embed="...">
    for drawing_el in para_element.iter(f"{{{_NSMAP_W}}}drawing"):
        for blip in drawing_el.iter(f"{{{_NSMAP_A}}}blip"):
            rid = blip.get(r_embed)
            if rid:
                rel_ids.append(rid)

    # 路径 2: VML — <w:pict> / <v:imagedata r:id="...">
    for pict_el in para_element.iter(f"{{{_NSMAP_W}}}pict"):
        for imgdata in pict_el.iter(f"{{{_NSMAP_V}}}imagedata"):
            rid = imgdata.get(r_id_attr)
            if rid:
                rel_ids.append(rid)

    return rel_ids


def _textbox_texts(para_element) -> List[str]:
    """收集段落/单元格元素内全部文本框（w:txbxContent）的文字。

    para.text / cell.text（python-docx）只拼接直属 run 的文字——锚定在段落里的
    文本框内容整段静默丢失。富岭 SOP 大量用文本框写步骤说明（FL-XS-WI-005
    实证：72 个文本框，"在电脑桌面打开U8/输入密码登录"等关键步骤全在框内，
    pdf 渲染可见而 docx 抽取为空，造成转换孪生"pdf 独有内容"假象 — 2026-06-12）。

    注意：
    - DrawingML（wps:txbx）与 VML 回退（v:textbox）各包一份相同的 w:txbxContent
      （mc:AlternateContent 的 Choice/Fallback 双份）→ 按文本去重；
    - 纯圈号/单字符标注框（"④"箭头标签类）是版面视觉标注，不入正文。
    """
    import re as _re
    texts: List[str] = []
    seen = set()
    for txbx in para_element.iter(f"{{{_NSMAP_W}}}txbxContent"):
        lines = []
        for p in txbx.iter(f"{{{_NSMAP_W}}}p"):
            line = "".join(t.text or "" for t in p.iter(f"{{{_NSMAP_W}}}t")).strip()
            if line:
                lines.append(line)
        s = "\n".join(lines).strip()
        if not s or s in seen:
            continue
        if _re.fullmatch(r"[①-⑳\s　]{1,2}", s):
            continue
        seen.add(s)
        texts.append(s)
    return texts


def _cell_text_with_textboxes(cell) -> str:
    """单元格文字 + 框内文字（cell.text 不含文本框）。"""
    cell_text = cell.text.strip()
    tb = _textbox_texts(cell._tc)
    if tb:
        cell_text = "\n".join(x for x in [cell_text] + tb if x).strip()
    return cell_text


def _has_images(para_element) -> bool:
    """快速判断 <w:p> 是否包含嵌入图片（不提取 rId，仅做存在性检查）。"""
    for _ in para_element.iter(f"{{{_NSMAP_W}}}drawing"):
        return True
    for _ in para_element.iter(f"{{{_NSMAP_W}}}pict"):
        return True
    return False


def extract_docx_with_images(
    local_path: str,
) -> Tuple[List[ExtractedBlock], List["ImageAsset"], List[str]]:
    """
    带图片位置追踪的 DOCX 提取（方案 B — heuristic）。

    在文档 body 的 XML 元素级遍历段落与表格，对每个 <w:p>：
      - 若仅含文本 → 发射 heading / paragraph / table 类型的 ExtractedBlock
      - 若含嵌入图片 → 发射 image_ref 类型的 ExtractedBlock（text="", extra 含 image_index/rel_id）
      - 若同时含文字和图片 → 先发射文字 block，再发射 image_ref block

    图片 ImageAsset 对象按文档顺序收集，其 image_index 与 ExtractedBlock.extra["image_index"] 对齐。

    注意：此函数不会将图片导出到磁盘（那是 extract_images_from_docx 的职责），
    它只记录 image_ref 占位 block 以标记图片在文档流中的位置。

    Args:
        local_path: DOCX 文件的本地路径。

    Returns:
        (blocks, image_assets_in_order)
        - blocks: 包含 text 和 image_ref 块的有序列表
        - image_assets_in_order: 按文档出现顺序排列的 ImageAsset 列表（local_path 为空串，
          仅记录 image_index 和 original_name/rel_id）
    """
    from opensearch_pipeline.extraction.image_extraction_utils import ImageAsset

    try:
        import docx
        from docx.text.paragraph import Paragraph
        from docx.table import Table
    except ImportError:
        return [], [], ["python-docx not installed, cannot extract DOCX"]

    try:
        document = docx.Document(local_path)
    except Exception as e:
        # 损坏 / 加密 / 截断下载的 DOCX：之前静默 return [], [] → 上游 warnings=[]，
        # 0 chunk 却以 DONE(成功) 收尾，坏文档无任何信号。回传 warning 让其可被
        # node_write_chunk_meta 的"疑似失败"判定捕获（DOCX 无 OCR 兜底，必须显式留痕）。
        return [], [], [f"Failed to open DOCX: {e}"]

    blocks: List[ExtractedBlock] = []
    image_assets: List[ImageAsset] = []
    warnings: List[str] = []
    current_section: Optional[str] = None
    image_counter = 0

    # 获取 document part 用于 rel_id → target_ref 映射
    doc_part = document.part

    def _rel_id_to_target_ref(rel_id: str) -> str:
        """尝试通过 rel_id 获取图片在包内的路径（如 media/image3.jpeg）。"""
        try:
            rel = doc_part.rels.get(rel_id)
            if rel and hasattr(rel, 'target_ref'):
                return rel.target_ref
        except Exception:
            pass
        return ""

    body = document.element.body

    for child in body.iterchildren():
        tag = child.tag

        # ── 段落 ──────────────────────────────────────────────
        if tag.endswith('}p') or tag == 'p':
            para = Paragraph(child, document)
            text = para.text.strip()
            style_name = para.style.name if para.style else "Normal"

            has_img = _has_images(child)

            # 先处理文本部分（如果有）
            if text:
                heading_level = _detect_heading_level(style_name, text)
                if heading_level is not None:
                    current_section = text
                    blocks.append(ExtractedBlock(
                        block_type="heading",
                        text=text,
                        level=heading_level,
                        section_path=current_section,
                        source="native",
                        extra={"word_style": style_name},
                    ))
                else:
                    blocks.append(ExtractedBlock(
                        block_type="paragraph",
                        text=text,
                        section_path=current_section,
                        source="native",
                        extra={"word_style": style_name},
                    ))

            # 文本框内容（para.text 不含，w:txbxContent）→ 独立 paragraph block
            for tb in _textbox_texts(child):
                blocks.append(ExtractedBlock(
                    block_type="paragraph",
                    text=tb,
                    section_path=current_section,
                    source="native",
                    extra={"word_style": style_name, "from_textbox": True},
                ))

            # 再处理图片引用（如果有）
            if has_img:
                rel_ids = _find_image_rel_ids(child)
                if not rel_ids:
                    # 检测到 drawing/pict 元素但无法解析 rel_id，仍记录占位
                    rel_ids = [""]

                for rid in rel_ids:
                    target_ref = _rel_id_to_target_ref(rid) if rid else ""

                    blocks.append(ExtractedBlock(
                        block_type="image_ref",
                        text="",
                        section_path=current_section,
                        source="native",
                        extra={
                            "image_index": image_counter,
                            "rel_id": rid,
                            "target_ref": target_ref,
                        },
                    ))

                    image_assets.append(ImageAsset(
                        local_path="",          # 尚未导出到磁盘
                        page_num=None,          # DOCX 无原生页码
                        image_index=image_counter,
                        original_name=target_ref,
                    ))
                    image_counter += 1

        # ── 表格 ──────────────────────────────────────────────
        elif tag.endswith('}tbl') or tag == 'tbl':
            table = Table(child, document)

            # 与 extract_docx 一致：解包 1×1 装饰性外表格
            if len(table.rows) == 1 and len(table.rows[0].cells) == 1:
                cell = table.rows[0].cells[0]
                # 递归处理单单元格内容：遍历子元素
                for sub_child in cell._tc.iterchildren():
                    sub_tag = sub_child.tag
                    if sub_tag.endswith('}p') or sub_tag == 'p':
                        sub_para = Paragraph(sub_child, cell)
                        sub_text = sub_para.text.strip()
                        sub_style = sub_para.style.name if sub_para.style else "Normal"
                        sub_has_img = _has_images(sub_child)

                        if sub_text:
                            sub_heading = _detect_heading_level(sub_style, sub_text)
                            if sub_heading is not None:
                                current_section = sub_text
                                blocks.append(ExtractedBlock(
                                    block_type="heading",
                                    text=sub_text,
                                    level=sub_heading,
                                    section_path=current_section,
                                    source="native",
                                    extra={"word_style": sub_style},
                                ))
                            else:
                                blocks.append(ExtractedBlock(
                                    block_type="paragraph",
                                    text=sub_text,
                                    section_path=current_section,
                                    source="native",
                                    extra={"word_style": sub_style},
                                ))

                        for tb in _textbox_texts(sub_child):
                            blocks.append(ExtractedBlock(
                                block_type="paragraph",
                                text=tb,
                                section_path=current_section,
                                source="native",
                                extra={"word_style": sub_style, "from_textbox": True},
                            ))

                        if sub_has_img:
                            sub_rel_ids = _find_image_rel_ids(sub_child)
                            if not sub_rel_ids:
                                sub_rel_ids = [""]
                            for rid in sub_rel_ids:
                                target_ref = _rel_id_to_target_ref(rid) if rid else ""
                                blocks.append(ExtractedBlock(
                                    block_type="image_ref",
                                    text="",
                                    section_path=current_section,
                                    source="native",
                                    extra={
                                        "image_index": image_counter,
                                        "rel_id": rid,
                                        "target_ref": target_ref,
                                    },
                                ))
                                image_assets.append(ImageAsset(
                                    local_path="",
                                    page_num=None,
                                    image_index=image_counter,
                                    original_name=target_ref,
                                ))
                                image_counter += 1
            else:
                rows_text = []
                # DC-3: dedup merged cells (gridSpan/vMerge) — see note in extract_docx above; store
                # the <w:tc> element itself (not id()) so lxml's per-node proxy keeps identity stable.
                seen_tc = set()
                for row in table.rows:
                    cells = []
                    for cell in row.cells:
                        tc = getattr(cell, "_tc", None)
                        if tc is not None:
                            if tc in seen_tc:
                                continue
                            seen_tc.add(tc)
                        cell_text = _cell_text_with_textboxes(cell)
                        if cell_text:
                            cells.append(cell_text)
                    if cells:
                        rows_text.append(" | ".join(cells))

                if rows_text:
                    table_md = "\n".join(f"| {row} |" for row in rows_text)
                    blocks.append(ExtractedBlock(
                        block_type="table",
                        text=table_md,
                        section_path=current_section,
                        source="native",
                        extra={"row_count": len(rows_text)},
                    ))

    return blocks, image_assets, warnings
