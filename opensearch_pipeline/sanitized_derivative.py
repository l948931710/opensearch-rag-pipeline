# -*- coding: utf-8 -*-
"""Controlled sanitized-derivative capability — codifies the 2026-06-20 quarantine-derivative flow.

When a quarantined doc is a procedure wrapped around an embedded PII roster, the safe
serving artifact is a **new, self-contained sanitized derivative**: the excised
procedure text + only the verified-clean step screenshots, ingested as a brand-new
source/doc/version that NEVER references the original canonical/assets. This module
turns that ad-hoc flow into reusable, tested primitives + an explicit safety contract.

HARD DEFAULT: **never auto-publish.** `DEFAULT_AUTO_PUBLISH = False`. There is no
function here that runs ingest end-to-end and serves the result. Prod execution stays
operator-gated per stage (source+register → Stage1 → rescan → Stage2 → Stage3), each
needing an explicit same-day PROD-RW authorization, and the final publish is gated by
`publish_authorized()` which requires ALL of: local PII rescan = 0 high, a Claude judge
approval (pii_residual=pass, semantic≥4, over_redaction=pass), AND a human spot-check.

Invariants the helpers enforce:
  - the package is self-contained (text + embedded clean images), images PIL-renormalized;
  - it carries NO original doc_id, asset/canonical/quarantine paths, roster image refs,
    redaction placeholders, or high-severity PII (`validate_package_no_residue`);
  - new independent source_key/doc_id/version (caller supplies; helpers never reuse the
    original).
"""
from __future__ import annotations

import hashlib
import io
import re
from typing import Any, Callable, Dict, List, Optional, Tuple

from opensearch_pipeline import redaction as R

SANITIZED_VERSION = "sanitized-derivative-v1"
DEFAULT_AUTO_PUBLISH = False  # load-bearing: this capability never publishes on its own.

# what publish() (operator-run, elsewhere) must have satisfied before serving a derivative.
PUBLISH_REQUIREMENTS = (
    "local_pii_high_zero",        # local text+image re-OCR high-severity PII == 0
    "judge_pii_residual_pass",    # Claude judge: pii_residual_risk == pass
    "judge_semantic_ge4",         # Claude judge: semantic_preservation >= 4
    "judge_over_redaction_pass",  # Claude judge: over_redaction == pass
    "human_spotcheck",            # mandatory human eyeball sign-off
)

_STEP_SPLIT = re.compile(r'(?m)^(?=\d、)')


def _sha(b: bytes) -> str:
    return hashlib.sha256(b).hexdigest()


def _normalize_jpeg(image_bytes: bytes, quality: int = 90) -> bytes:
    """Re-encode to a clean baseline JPEG (python-docx rejects some CMYK/progressive headers)."""
    from PIL import Image
    im = Image.open(io.BytesIO(image_bytes))
    if im.mode != "RGB":
        im = im.convert("RGB")
    buf = io.BytesIO()
    im.save(buf, format="JPEG", quality=quality)
    return buf.getvalue()


def build_sanitized_source_docx(
    title: str,
    procedure_text: str,
    images: Optional[List[Dict[str, Any]]] = None,
) -> Dict[str, Any]:
    """Build a self-contained sanitized DOCX (procedure text + step-bound images).

    Args:
      images: [{"image_bytes": bytes, "bind_step": "3"}] — each embedded right after its
              step paragraph so the docx extractor's positional refs bind it deterministically.
    Returns: {docx_bytes, pkg_sha256, procedure_text_sha256, image_count, images:[{bind_step,
              source_sha256, embedded_sha256}]}.
    """
    import docx
    from docx.shared import Inches
    images = images or []
    metas, emb_by_step = [], {}
    for img in images:
        raw = img["image_bytes"]
        cj = _normalize_jpeg(raw)
        step = str(img.get("bind_step") or "")
        emb_by_step.setdefault(step, []).append(cj)
        metas.append({"bind_step": step, "source_sha256": _sha(raw), "embedded_sha256": _sha(cj)})

    d = docx.Document()
    d.add_heading(title, level=1)
    for seg in [s.strip() for s in _STEP_SPLIT.split(procedure_text) if s.strip()]:
        d.add_paragraph(seg)
        m = re.match(r'^(\d)、', seg)
        if m:
            for cj in emb_by_step.get(m.group(1), []):
                d.add_picture(io.BytesIO(cj), width=Inches(5.2))
    buf = io.BytesIO()
    d.save(buf)
    pkg = buf.getvalue()
    return {
        "docx_bytes": pkg,
        "pkg_sha256": _sha(pkg),
        "procedure_text_sha256": _sha(procedure_text.encode("utf-8")),
        "image_count": len(metas),
        "images": metas,
    }


# patterns that must NOT appear in a sanitized source package body
_ASSET_PATH = re.compile(r'processing/(assets|canonical)|/_quarantine/|raw/[^/]+/_quarantine/')
_ROSTER_IMGREF = re.compile(r'\b(p\d+_img\d+|image_refs|oss_key)\b')
_PLACEHOLDER = re.compile(r'已脱敏\]')


def validate_package_no_residue(
    docx_bytes: bytes,
    *,
    forbidden_doc_ids: Tuple[str, ...] = (),
    name_llm_fn: Optional[Callable[[str], str]] = None,
) -> Dict[str, Any]:
    """Re-open the built docx and prove it carries no original-doc references or PII.

    Checks: no forbidden doc_id; no asset/canonical/quarantine path; no roster image-ref
    token; no redaction placeholder; no high-severity PII (id/partial-id/mobile/bank/
    address/name via the redaction engine). 办件单号-style 19-22 digit case numbers are
    reported separately (institutional, non-personal — informational, not a violation).
    """
    import docx
    d = docx.Document(io.BytesIO(docx_bytes))
    body = " ".join(p.text for p in d.paragraphs)
    n_imgs = sum(1 for r in d.part.rels.values() if "image" in r.reltype)
    v = {}
    v["forbidden_doc_id"] = sorted(x for x in forbidden_doc_ids if x and x in body)
    v["asset_path"] = bool(_ASSET_PATH.search(body))
    v["roster_image_ref"] = bool(_ROSTER_IMGREF.search(body))
    v["redaction_placeholder"] = bool(_PLACEHOLDER.search(body))
    v["high_pii"] = R.high_residual(R.rescan(body, name_llm_fn=name_llm_fn))
    violations = (v["forbidden_doc_id"] or v["asset_path"] or v["roster_image_ref"]
                  or v["redaction_placeholder"] or bool(v["high_pii"]))
    return {
        "ok": not violations,
        "violations": v,
        "image_count": n_imgs,
        "text_chars": len(body),
        "case_numbers_present": bool(re.search(r'(?<!\d)\d{19,22}(?!\d)', body)),  # info-only (institutional)
    }


def publish_authorized(gates: Dict[str, Any]) -> Tuple[bool, List[str]]:
    """Gate for serving a sanitized derivative. ALL PUBLISH_REQUIREMENTS must be truthy.

    Default-deny: a missing/falsey requirement blocks publish. Returns (ok, missing[]).
    """
    missing = [req for req in PUBLISH_REQUIREMENTS if not gates.get(req)]
    return (not missing), missing


# Documented gated phases — prod execution is operator-driven, one PROD-RW-authorized
# step at a time. This is reference/structure only; it does not execute prod writes.
GATED_PHASES = (
    "BUILD_SOURCE_PACKAGE",   # build_sanitized_source_docx + validate_package_no_residue
    "REGISTER",               # scoped register: new independent doc_id/version/source_key
    "STAGE1_CANONICALIZE",    # extract → canonical (+ full PII rescan gate)
    "STAGE2_CHUNK",           # chunk + step-card binding (+ chunk-level PII/perm/structure gate)
    "STAGE3_INDEX",           # embed + HA3 push (+ RDS↔HA3 self-query parity + ACL matrix)
    "PUBLISH_DECISION",       # publish_authorized() must pass (incl. human spot-check)
)
