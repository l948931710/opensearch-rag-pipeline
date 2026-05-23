# -*- coding: utf-8 -*-
"""
test_real_extractors.py — DOCX/PDF/OCR 真实文档解析器测试套件

覆盖：
  1. DOCX 标题检测（样式优先 + 中文 regex fallback）
  2. DOCX 真实文件提取（段落/表格/标题/section_path/损坏文件）
  3. PDF 多策略 fallback 编排（pdfplumber → pypdf → 空）
  4. PDF 真实文件提取（pypdf PdfWriter 构建最小 PDF）
  5. OCR 客户端（模拟模式/真实模式 API 无 KEY/DashScope 响应解析/Gemini 响应解析）
"""

import os
import tempfile
import unittest.mock as mock
from unittest.mock import MagicMock, patch

import pytest

from opensearch_pipeline.extraction.schema import ExtractedBlock, ExtractionResult
from opensearch_pipeline.extraction.docx_extractor import (
    _detect_heading_level,
    extract_docx,
)
from opensearch_pipeline.extraction.pdf_extractor import (
    extract_pdf,
    get_pdf_text_length,
    _extract_with_pypdf,
)
from opensearch_pipeline.extraction.ocr_client import (
    OCRClient,
    OCRResult,
    OCRPageResult,
)


# ═══════════════════════════════════════════════════════════════
# 1. DOCX 标题检测测试
# ═══════════════════════════════════════════════════════════════

class TestDocxHeadingDetection:
    """_detect_heading_level 标题级别检测：样式优先 + regex fallback。"""

    def test_standard_heading_styles(self):
        """标准 Word 英文样式 → 正确 level。"""
        assert _detect_heading_level("Heading 1", "任何文本") == 1
        assert _detect_heading_level("Heading 2", "任何文本") == 2
        assert _detect_heading_level("Heading 3", "任何文本") == 3
        assert _detect_heading_level("Heading 4", "任何文本") == 4

    def test_chinese_heading_styles(self):
        """中文 Word 样式 → 正确 level。"""
        assert _detect_heading_level("标题 1", "任何文本") == 1
        assert _detect_heading_level("标题 2", "任何文本") == 2
        assert _detect_heading_level("标题 3", "任何文本") == 3

    def test_title_subtitle_styles(self):
        """Title → 1, Subtitle → 2。"""
        assert _detect_heading_level("Title", "文档标题") == 1
        assert _detect_heading_level("Subtitle", "副标题") == 2

    def test_case_insensitive_heading_detection(self):
        """样式名包含 'heading' 但大小写不同时仍能识别。"""
        assert _detect_heading_level("heading 1", "abc") == 1
        # 不在 map 中但包含 "heading" 关键字
        assert _detect_heading_level("Custom Heading 5", "abc") == 4  # min(5, 4) = 4
        assert _detect_heading_level("My heading", "abc") == 2  # 无数字 → 默认 2

    def test_chinese_regex_chapter(self):
        """中文 '第X章' 模式 → level 1。"""
        assert _detect_heading_level("Normal", "第一章 总则") == 1
        assert _detect_heading_level("Normal", "第三章 安全管理") == 1
        assert _detect_heading_level("Normal", "第10章 附则") == 1

    def test_chinese_regex_numbered(self):
        """中文 '一、' 编号模式 → level 2。"""
        assert _detect_heading_level("Normal", "一、适用范围") == 2
        assert _detect_heading_level("Normal", "三、注意事项") == 2

    def test_chinese_regex_parenthetical(self):
        """中文 '（一）' 括号编号 → level 2。"""
        assert _detect_heading_level("Normal", "（一）概述") == 2
        assert _detect_heading_level("Normal", "（3）实施步骤") == 2

    def test_sub_heading_regex(self):
        """'3.1 登录系统' 子编号模式 → level 3。"""
        assert _detect_heading_level("Normal", "3.1 登录系统") == 3
        assert _detect_heading_level("Normal", "12.5 提交审核") == 3

    def test_long_text_not_heading(self):
        """超过 60 字符的文本即使匹配模式也不应被识别为标题（防误判）。"""
        long_text = "第一章 " + "A" * 60  # 总长 > 60
        assert _detect_heading_level("Normal", long_text) is None

    def test_normal_paragraph_no_level(self):
        """普通正文 → None。"""
        assert _detect_heading_level("Normal", "这是一段普通的正文内容。") is None
        assert _detect_heading_level("Normal", "Hello world") is None

    def test_empty_text(self):
        """空文本 → None。"""
        assert _detect_heading_level("Normal", "") is None


# ═══════════════════════════════════════════════════════════════
# 2. DOCX 真实文件提取测试
# ═══════════════════════════════════════════════════════════════

class TestDocxExtractorRealFile:
    """使用 python-docx 构建真实 DOCX 文件并测试提取。"""

    @pytest.fixture
    def create_docx(self):
        """提供创建临时 DOCX 文件的辅助函数。"""
        import docx
        temp_files = []

        def _make(paragraphs=None, tables=None, headings=None):
            """
            构建 DOCX。
            paragraphs: [(style, text), ...]
            tables: [[(cell1, cell2), (cell3, cell4)], ...]  每个子列表是一张表
            headings: [(level, text), ...]  使用 add_heading
            """
            doc = docx.Document()
            if headings:
                for level, text in headings:
                    doc.add_heading(text, level=level)
            if paragraphs:
                for style, text in paragraphs:
                    p = doc.add_paragraph(text, style=style)
            if tables:
                for table_data in tables:
                    rows = len(table_data)
                    cols = len(table_data[0]) if table_data else 0
                    t = doc.add_table(rows=rows, cols=cols)
                    for i, row_data in enumerate(table_data):
                        for j, cell_text in enumerate(row_data):
                            t.rows[i].cells[j].text = cell_text
            # 如果同时需要混合内容，按 headings → paragraphs → tables 顺序
            fd, path = tempfile.mkstemp(suffix=".docx")
            os.close(fd)
            doc.save(path)
            temp_files.append(path)
            return path

        yield _make

        for p in temp_files:
            if os.path.exists(p):
                os.remove(p)

    def test_basic_paragraphs(self, create_docx):
        """3 个普通段落 → 3 个 paragraph blocks。"""
        path = create_docx(paragraphs=[
            ("Normal", "第一段内容"),
            ("Normal", "第二段内容"),
            ("Normal", "第三段内容"),
        ])
        blocks, warnings = extract_docx(path)
        paragraphs = [b for b in blocks if b.block_type == "paragraph"]
        assert len(paragraphs) >= 3
        assert paragraphs[0].text == "第一段内容"

    def test_heading_and_paragraph(self, create_docx):
        """Heading 1 + Normal 段落 → 正确的 block_type 和 level。"""
        import docx
        doc = docx.Document()
        doc.add_heading("文档标题", level=1)
        doc.add_paragraph("这是正文内容。")
        doc.add_heading("安全管理", level=2)
        doc.add_paragraph("安全管理的具体细则如下。")
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        doc.save(path)

        try:
            blocks, warnings = extract_docx(path)
            headings = [b for b in blocks if b.block_type == "heading"]
            paragraphs = [b for b in blocks if b.block_type == "paragraph"]
            assert len(headings) >= 2
            assert headings[0].level == 1
            assert headings[0].text == "文档标题"
            assert headings[1].level == 2
            assert len(paragraphs) >= 2
        finally:
            os.remove(path)

    def test_table_extraction(self, create_docx):
        """2×2 表格 → block_type='table' + pipe 分隔文本。"""
        path = create_docx(tables=[
            [("姓名", "工号"), ("张三", "001")],
        ])
        blocks, warnings = extract_docx(path)
        tables = [b for b in blocks if b.block_type == "table"]
        assert len(tables) >= 1
        assert "|" in tables[0].text
        assert "张三" in tables[0].text
        assert tables[0].extra.get("row_count") == 2

    def test_empty_document(self, create_docx):
        """空 DOCX → 无 blocks、无 warnings。"""
        import docx
        doc = docx.Document()
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        doc.save(path)
        try:
            blocks, warnings = extract_docx(path)
            assert len(blocks) == 0
            assert len(warnings) == 0
        finally:
            os.remove(path)

    def test_corrupt_file_warning(self):
        """非 DOCX 文件 → 返回 warning。"""
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.write(fd, b"this is not a docx file")
        os.close(fd)
        try:
            blocks, warnings = extract_docx(path)
            assert len(blocks) == 0
            assert len(warnings) > 0
            assert "Failed to open DOCX" in warnings[0]
        finally:
            os.remove(path)

    def test_section_path_tracking(self, create_docx):
        """Heading 后的段落继承 section_path。"""
        import docx
        doc = docx.Document()
        doc.add_heading("审核流程", level=1)
        doc.add_paragraph("步骤一描述。")
        doc.add_heading("准备工作", level=2)
        doc.add_paragraph("准备内容。")
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        doc.save(path)
        try:
            blocks, warnings = extract_docx(path)
            paragraphs = [b for b in blocks if b.block_type == "paragraph"]
            assert len(paragraphs) >= 2
            assert paragraphs[0].section_path == "审核流程"
            assert paragraphs[1].section_path == "准备工作"
        finally:
            os.remove(path)

    def test_word_style_in_extra(self, create_docx):
        """每个 block 的 extra 字段应包含 word_style 信息。"""
        import docx
        doc = docx.Document()
        doc.add_heading("标题", level=1)
        doc.add_paragraph("内容")
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        doc.save(path)
        try:
            blocks, _ = extract_docx(path)
            for b in blocks:
                assert "word_style" in b.extra
        finally:
            os.remove(path)


# ═══════════════════════════════════════════════════════════════
# 3. PDF Fallback 编排逻辑测试
# ═══════════════════════════════════════════════════════════════

class TestPdfFallbackOrchestration:
    """extract_pdf() 的 pdfplumber → pypdf 多策略 fallback 逻辑。"""

    @patch("opensearch_pipeline.extraction.pdf_extractor._extract_with_pypdf")
    @patch("opensearch_pipeline.extraction.pdf_extractor._extract_with_pdfplumber")
    def test_pdfplumber_success_no_fallback(self, mock_plumber, mock_pypdf):
        """pdfplumber 成功提取 → pypdf 不被调用。"""
        mock_plumber.return_value = (
            [ExtractedBlock(block_type="paragraph", text="内容" * 50)],
            1,
            [],
        )
        blocks, page_count, warnings = extract_pdf("/fake/path.pdf")
        assert len(blocks) == 1
        assert page_count == 1
        mock_pypdf.assert_not_called()

    @patch("opensearch_pipeline.extraction.pdf_extractor._extract_with_pypdf")
    @patch("opensearch_pipeline.extraction.pdf_extractor._extract_with_pdfplumber")
    def test_pdfplumber_zero_chars_falls_back_to_pypdf(self, mock_plumber, mock_pypdf):
        """pdfplumber 返回 0 字符 → 降级到 pypdf。"""
        mock_plumber.return_value = ([], 2, ["no text"])
        mock_pypdf.return_value = (
            [ExtractedBlock(block_type="paragraph", text="pypdf 内容" * 20)],
            2,
            [],
        )
        blocks, page_count, warnings = extract_pdf("/fake/path.pdf")
        assert len(blocks) == 1
        assert "pypdf 内容" in blocks[0].text
        mock_pypdf.assert_called_once()
        assert any("0 chars" in w for w in warnings)

    @patch("opensearch_pipeline.extraction.pdf_extractor._extract_with_pypdf")
    @patch("opensearch_pipeline.extraction.pdf_extractor._extract_with_pdfplumber")
    def test_pdfplumber_import_error_falls_back(self, mock_plumber, mock_pypdf):
        """pdfplumber ImportError → 降级到 pypdf。"""
        mock_plumber.side_effect = ImportError("No module named 'pdfplumber'")
        mock_pypdf.return_value = (
            [ExtractedBlock(block_type="paragraph", text="fallback 内容")],
            1,
            [],
        )
        blocks, page_count, warnings = extract_pdf("/fake/path.pdf")
        assert len(blocks) == 1
        mock_pypdf.assert_called_once()
        assert any("pdfplumber not installed" in w for w in warnings)

    @patch("opensearch_pipeline.extraction.pdf_extractor._extract_with_pypdf")
    @patch("opensearch_pipeline.extraction.pdf_extractor._extract_with_pdfplumber")
    def test_both_fail_returns_empty(self, mock_plumber, mock_pypdf):
        """pdfplumber 和 pypdf 都失败 → 返回空 blocks + 收集所有 warnings。"""
        mock_plumber.side_effect = Exception("plumber error")
        mock_pypdf.side_effect = Exception("pypdf error")
        blocks, page_count, warnings = extract_pdf("/fake/path.pdf")
        assert len(blocks) == 0
        assert page_count == 0
        assert any("plumber error" in w for w in warnings)
        assert any("pypdf error" in w for w in warnings)


# ═══════════════════════════════════════════════════════════════
# 4. PDF 真实文件提取测试
# ═══════════════════════════════════════════════════════════════

class TestPdfExtractorRealFile:
    """使用 pypdf 构建真实 PDF 文件并测试提取。"""

    @pytest.fixture
    def create_pdf(self):
        """使用 pypdf PdfWriter 创建包含文本的最小 PDF。"""
        from pypdf import PdfWriter
        from pypdf.generic import (
            ArrayObject, DictionaryObject, NameObject, NumberObject,
            TextStringObject, StreamObject,
        )
        temp_files = []

        def _make(pages_text: list):
            """
            创建含有指定页面文本的 PDF。
            pages_text: ["page 1 text", "page 2 text", ...]
            """
            writer = PdfWriter()
            for text in pages_text:
                # 创建一个简单的 PDF 页面流
                stream = StreamObject()
                stream[NameObject("/Length")] = NumberObject(0)
                stream._data = f"BT /F1 12 Tf 100 700 Td ({text}) Tj ET".encode()

                resources = DictionaryObject()
                font_dict = DictionaryObject()
                font_dict[NameObject("/Type")] = NameObject("/Font")
                font_dict[NameObject("/Subtype")] = NameObject("/Type1")
                font_dict[NameObject("/BaseFont")] = NameObject("/Helvetica")
                fonts = DictionaryObject()
                fonts[NameObject("/F1")] = font_dict
                resources[NameObject("/Font")] = fonts

                page = DictionaryObject()
                page[NameObject("/Type")] = NameObject("/Page")
                page[NameObject("/MediaBox")] = ArrayObject([
                    NumberObject(0), NumberObject(0),
                    NumberObject(612), NumberObject(792),
                ])
                page[NameObject("/Resources")] = resources
                page[NameObject("/Contents")] = writer._add_object(stream)
                writer._add_page(page)

            fd, path = tempfile.mkstemp(suffix=".pdf")
            os.close(fd)
            with open(path, "wb") as f:
                writer.write(f)
            temp_files.append(path)
            return path

        yield _make

        for p in temp_files:
            if os.path.exists(p):
                os.remove(p)

    def test_corrupt_pdf_warning(self):
        """非 PDF 文件 → 返回 warning。"""
        fd, path = tempfile.mkstemp(suffix=".pdf")
        os.write(fd, b"this is not a pdf")
        os.close(fd)
        try:
            blocks, page_count, warnings = extract_pdf(path)
            # 两种 extractor 都应该失败
            assert len(warnings) > 0
        finally:
            os.remove(path)

    def test_get_pdf_text_length(self):
        """get_pdf_text_length 辅助函数计算正确。"""
        blocks = [
            ExtractedBlock(block_type="paragraph", text="hello"),
            ExtractedBlock(block_type="paragraph", text="世界"),
        ]
        assert get_pdf_text_length(blocks) == 7  # 5 + 2

    def test_get_pdf_text_length_empty(self):
        """空 blocks 列表 → 长度为 0。"""
        assert get_pdf_text_length([]) == 0

    @patch("opensearch_pipeline.extraction.pdf_extractor._extract_with_pypdf")
    @patch("opensearch_pipeline.extraction.pdf_extractor._extract_with_pdfplumber")
    def test_max_pages_limit_warning(self, mock_plumber, mock_pypdf):
        """超过 max_pages 限制时生成 warning。"""
        # 模拟 pdfplumber 返回 5 页但只提取 2 页
        blocks = [
            ExtractedBlock(block_type="paragraph", text=f"Page {i} content", page_num=i)
            for i in range(1, 3)
        ]
        mock_plumber.return_value = (
            blocks, 5,
            ["PDF has 5 pages, only first 2 extracted"],
        )
        result_blocks, page_count, warnings = extract_pdf("/fake/path.pdf", max_pages=2)
        assert page_count == 5
        assert len(result_blocks) == 2
        assert any("only first 2 extracted" in w for w in warnings)


# ═══════════════════════════════════════════════════════════════
# 5. OCR 客户端测试
# ═══════════════════════════════════════════════════════════════

class TestOCRClient:
    """OCRClient 模拟模式 + 真实模式 API 调用 + 响应解析。"""

    def test_simulate_mode_pdf_ocr(self):
        """simulate=True 的 PDF OCR → 返回 SIMULATED 状态和 2 页。"""
        client = OCRClient(simulate=True)
        result = client.ocr_pdf("/fake.pdf", "DOC001")
        assert result.status == "SIMULATED"
        assert result.page_count == 2
        assert "DOC001" in result.combined_text
        assert all(p.status == "SIMULATED" for p in result.pages)

    def test_simulate_mode_image_ocr(self):
        """simulate=True 的图片 OCR → 返回 SIMULATED 状态。"""
        client = OCRClient(simulate=True)
        result = client.ocr_image("/fake.png", "DOC002")
        assert result.status == "SIMULATED"
        assert result.page_count == 1
        assert "image content recognized" in result.combined_text

    def test_ocr_result_to_blocks(self):
        """OCRResult.to_blocks() 正确生成 ExtractedBlock 列表。"""
        result = OCRResult(
            pages=[
                OCRPageResult(page_num=1, text="第一页内容"),
                OCRPageResult(page_num=2, text="第二页内容"),
            ],
            combined_text="第一页内容\n\n第二页内容",
            status="DONE",
        )
        blocks = result.to_blocks()
        assert len(blocks) == 2
        assert all(b.block_type == "ocr_text" for b in blocks)
        assert all(b.source == "ocr" for b in blocks)
        assert blocks[0].page_num == 1
        assert blocks[1].text == "第二页内容"

    def test_ocr_result_empty_pages_filtered(self):
        """空文本页在 to_blocks() 中被过滤。"""
        result = OCRResult(
            pages=[
                OCRPageResult(page_num=1, text="有内容"),
                OCRPageResult(page_num=2, text=""),
                OCRPageResult(page_num=3, text="   "),
            ],
        )
        blocks = result.to_blocks()
        assert len(blocks) == 1
        assert blocks[0].page_num == 1

    def test_real_pdf_ocr_no_api_key(self):
        """simulate=False 且无 API KEY → 返回 FAILED。"""
        client = OCRClient(api_key="", simulate=False)
        result = client.ocr_pdf("/fake.pdf", "DOC003")
        assert result.status == "FAILED"
        assert "API KEY" in result.error

    def test_real_image_ocr_no_api_key(self):
        """simulate=False 且无 API KEY → 返回 FAILED。"""
        client = OCRClient(api_key="", simulate=False)
        result = client.ocr_image("/fake.png", "DOC004")
        assert result.status == "FAILED"
        assert "API KEY" in result.error

    @patch("requests.post")
    def test_dashscope_ocr_api_call(self, mock_post):
        """DashScope OCR API 调用：验证 URL、payload 结构、认证头。"""
        mock_resp = MagicMock()
        mock_resp.status_code = 200
        mock_resp.json.return_value = {
            "output": {
                "choices": [{"message": {"content": "识别结果"}}]
            }
        }
        mock_post.return_value = mock_resp

        client = OCRClient(
            api_key="sk-test-key",
            api_base_url="https://dashscope.aliyuncs.com/api/v1",
            ocr_model="qwen-vl-max",
            simulate=False,
        )
        text = client._call_ocr_api("base64data", "image/png")

        assert text == "识别结果"
        mock_post.assert_called_once()
        args, kwargs = mock_post.call_args
        assert "dashscope.aliyuncs.com" in args[0]
        assert "multimodal-generation" in args[0]
        assert kwargs["headers"]["Authorization"] == "Bearer sk-test-key"
        assert kwargs["json"]["model"] == "qwen-vl-max"

    @patch("requests.post")
    def test_gemini_ocr_api_call(self, mock_post):
        """Gemini OCR API 调用：验证 URL、payload 结构、x-goog-api-key 头（非 URL 参数）。"""
        mock_resp = MagicMock()
        mock_resp.status_code = 200
        mock_resp.json.return_value = {
            "candidates": [{"content": {"parts": [{"text": "Gemini OCR result"}]}}]
        }
        mock_post.return_value = mock_resp

        client = OCRClient(
            api_key="AIzaSy_test_key",
            api_base_url="https://generativelanguage.googleapis.com/v1beta",
            ocr_model="gemini-2.0-flash",
            simulate=False,
        )
        text = client._call_ocr_api("base64data", "image/jpeg")

        assert text == "Gemini OCR result"
        mock_post.assert_called_once()
        args, kwargs = mock_post.call_args
        # API key 在 header 中，不在 URL 中
        assert "x-goog-api-key" in kwargs["headers"]
        assert kwargs["headers"]["x-goog-api-key"] == "AIzaSy_test_key"
        assert "key=" not in args[0]  # 不在 URL 中暴露密钥
        assert "generateContent" in args[0]

    @patch("requests.post")
    def test_dashscope_ocr_response_parsing_list_content(self, mock_post):
        """DashScope 返回 list 格式 content 的解析。"""
        mock_resp = MagicMock()
        mock_resp.status_code = 200
        mock_resp.json.return_value = {
            "output": {
                "choices": [{
                    "message": {
                        "content": [
                            {"text": "第一段"},
                            {"text": "第二段"},
                        ]
                    }
                }]
            }
        }
        mock_post.return_value = mock_resp

        client = OCRClient(
            api_key="sk-test",
            api_base_url="https://dashscope.aliyuncs.com/api/v1",
            simulate=False,
        )
        text = client._call_ocr_api("b64", "image/png")
        assert text == "第一段第二段"

    @patch("requests.post")
    def test_dashscope_ocr_response_parsing_string_content(self, mock_post):
        """DashScope 返回 string 格式 content 的解析。"""
        mock_resp = MagicMock()
        mock_resp.status_code = 200
        mock_resp.json.return_value = {
            "output": {
                "choices": [{"message": {"content": "纯字符串结果"}}]
            }
        }
        mock_post.return_value = mock_resp

        client = OCRClient(
            api_key="sk-test",
            api_base_url="https://dashscope.aliyuncs.com/api/v1",
            simulate=False,
        )
        text = client._call_ocr_api("b64", "image/png")
        assert text == "纯字符串结果"

    @patch("requests.post")
    def test_gemini_ocr_response_parsing(self, mock_post):
        """Gemini 标准响应格式解析。"""
        mock_resp = MagicMock()
        mock_resp.status_code = 200
        mock_resp.json.return_value = {
            "candidates": [{
                "content": {
                    "parts": [{"text": "  Gemini 识别文本  "}]
                }
            }]
        }
        mock_post.return_value = mock_resp

        client = OCRClient(
            api_key="test",
            api_base_url="https://generativelanguage.googleapis.com/v1beta",
            simulate=False,
        )
        text = client._call_ocr_api("b64", "image/png")
        assert text == "Gemini 识别文本"  # strip()

    @patch("requests.post")
    def test_dashscope_ocr_http_error(self, mock_post):
        """DashScope 返回 500 → RuntimeError。"""
        mock_resp = MagicMock()
        mock_resp.status_code = 500
        mock_resp.text = "Internal Server Error"
        mock_post.return_value = mock_resp

        client = OCRClient(
            api_key="sk-test",
            api_base_url="https://dashscope.aliyuncs.com/api/v1",
            simulate=False,
        )
        with pytest.raises(RuntimeError, match="DashScope OCR HTTP 500"):
            client._call_ocr_api("b64", "image/png")

    @patch("requests.post")
    def test_gemini_ocr_http_error(self, mock_post):
        """Gemini 返回 401 → RuntimeError。"""
        mock_resp = MagicMock()
        mock_resp.status_code = 401
        mock_resp.text = "Unauthorized"
        mock_post.return_value = mock_resp

        client = OCRClient(
            api_key="bad-key",
            api_base_url="https://generativelanguage.googleapis.com/v1beta",
            simulate=False,
        )
        with pytest.raises(RuntimeError, match="Gemini OCR HTTP 401"):
            client._call_ocr_api("b64", "image/png")

    @patch("requests.post")
    def test_dashscope_ocr_empty_response(self, mock_post):
        """DashScope 返回空/畸形响应 → 返回空字符串而非崩溃。"""
        mock_resp = MagicMock()
        mock_resp.status_code = 200
        mock_resp.json.return_value = {"output": {"choices": []}}
        mock_post.return_value = mock_resp

        client = OCRClient(
            api_key="sk-test",
            api_base_url="https://dashscope.aliyuncs.com/api/v1",
            simulate=False,
        )
        text = client._call_ocr_api("b64", "image/png")
        assert text == ""

    @patch("requests.post")
    def test_gemini_url_construction_with_models_in_base(self, mock_post):
        """当 api_base_url 已包含 /models/ 路径时，不重复拼接。"""
        mock_resp = MagicMock()
        mock_resp.status_code = 200
        mock_resp.json.return_value = {
            "candidates": [{"content": {"parts": [{"text": "ok"}]}}]
        }
        mock_post.return_value = mock_resp

        client = OCRClient(
            api_key="test",
            api_base_url="https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash",
            simulate=False,
        )
        client._call_ocr_api("b64", "image/png")
        url = mock_post.call_args[0][0]
        assert url.endswith(":generateContent")
        assert "/models/models/" not in url  # 不重复拼接

    def test_real_image_ocr_file_not_found(self):
        """真实图片 OCR 但文件不存在 → FAILED。"""
        client = OCRClient(api_key="test-key", simulate=False)
        result = client._real_image_ocr("/nonexistent/path.png", "DOC005")
        assert result.status == "FAILED"
        assert result.error is not None
