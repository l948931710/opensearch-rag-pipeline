# -*- coding: utf-8 -*-
"""Generate the committed binary extraction fixtures (run once; the binaries are committed to git).

These are small, deterministic, NON-sensitive files used to regression-test the real extractors in CI
without depending on out-of-repo prod data or synthesizing at test time. The docx carries merged table
cells (gridSpan + vMerge) as a permanent guard for the DC-3 dedup fix.

Regenerate:  python tests/fixtures/make_fixtures.py
"""
import os

HERE = os.path.dirname(os.path.abspath(__file__))


def make_docx():
    import docx
    d = docx.Document()
    d.add_heading("检验规范", level=1)
    d.add_paragraph("正文段落示例。")
    t = d.add_table(rows=2, cols=3)
    t.cell(0, 0).merge(t.cell(0, 1)).text = "合并表头"   # horizontal gridSpan (merged across cols 0-1)
    t.cell(0, 2).text = "列C"
    t.cell(1, 0).text = "甲"
    t.cell(1, 1).text = "乙"
    t.cell(1, 2).text = "丙"
    t2 = d.add_table(rows=2, cols=2)
    t2.cell(0, 0).merge(t2.cell(1, 0)).text = "纵向合并"  # vertical vMerge (merged across rows 0-1)
    t2.cell(0, 1).text = "P"
    t2.cell(1, 1).text = "Q"
    d.save(os.path.join(HERE, "merged_cells.docx"))


def make_xlsx():
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "检验表"
    ws["A1"] = "项目"
    ws["B1"] = "值"
    ws["A2"] = "温度"
    ws["B2"] = "25"
    ws.merge_cells("A3:B3")
    ws["A3"] = "合并行说明"
    wb.save(os.path.join(HERE, "merged_cells.xlsx"))


def make_pdf():
    import fitz
    doc = fitz.open()
    page = doc.new_page()
    # ASCII only — fitz's default font has no CJK glyphs; the PDF test only checks text extraction.
    page.insert_text((72, 72), "PDF fixture line one.", fontsize=12)
    page.insert_text((72, 96), "PDF fixture line two.", fontsize=12)
    doc.save(os.path.join(HERE, "sample.pdf"))
    doc.close()


def make_two_column_header_straddle_pdf():
    """双栏 + 页眉跨裁剪线的最小复现件（PDF-D2，2026-07-25）。

    复现 FL-XS-WI-007 p2 的三要素——三者缺一都不会触发图↔步骤错绑：
      1. 版面被 _detect_column_split 判为双栏（阅读序 = 左栏整栏 → 右栏整栏）；
      2. 页眉最后一行落在 10% 候选带**之外**（本件 top≈92 > 84.19）⇒ 不进页眉候选，
         但它又跨在 crop_top(=95) 上 ⇒ pdfplumber 裁剪保留该词并把 top **钳到 95**；
      3. 左栏底部有个步骤行（top≈497）。
    于是"左栏末行(497) → 右栏首行(钳位 95)"产生**负** gap，单向 `>40` 判据看不见，
    Step 4 的段落块 y 包络膨胀成 [95, 507] 罩住整页。

    坐标是被测断言的一部分（tests/test_pdf_column_gap_binding.py 冻结了实测值），
    改动这里必须同步更新那些常量。ASCII only：fitz 默认字体无中文字形，且
    STEP_BOUNDARY_PATTERN 支持 `Step N`。
    """
    import fitz
    W, H, FS, GAP = 595.276, 841.89, 10.0, 5.0
    LX, RX = 42.5, 400.0        # 左/右栏起点：中部 30%-70%（178.6..416.7）无词横跨

    def line(page, x, y, text):
        # 逐词显式排 x —— pdfplumber 的 x_tolerance=3 会吞掉普通空格、把整行粘成一个词
        for w in text.split():
            page.insert_text((x, y), w, fontsize=FS)
            x += fitz.get_text_length(w, fontname="helv", fontsize=FS) + GAP

    doc = fitz.open()
    for _ in range(2):          # 两页：页眉频次阈值 max(2, int(n*0.6)) 要求出现在两页
        p = doc.new_page(width=W, height=H)
        line(p, RX, 58,  "Fuling Tech Co Ltd")           # top≈50  → 页眉候选
        line(p, RX, 73,  "Doc No FL-TEST-001")           # top≈65  → 页眉候选
        line(p, RX, 88,  "Doc Version A0")               # top≈80  → 页眉候选（max ⇒ crop_top=80+15）
        line(p, RX, 100, "Effective Date 2022-08-19")    # top≈92  → 带外漏网 + 跨 crop_top
        line(p, LX, 131, "Step 3 open the U8 scan screen")
        line(p, LX, 146, "follow the marked buttons")
        # 下面两行纵向间距 50pt（>40）且非 step/heading/圈号 —— 正向 gap 分支的专用样本
        line(p, LX, 200, "the operator confirms the label")
        line(p, LX, 260, "then the quantity is checked")
        line(p, LX, 505, "Step 4 aim the scanner at the code")
        for i, t in enumerate(["click here", "open the app", "double click",
                               "enter code", "confirm now", "close window"]):
            line(p, RX, 160 + i * 55, t)   # 右栏词量与纵向重叠：喂饱 _detect_column_split
    doc.save(os.path.join(HERE, "two_column_header_straddle.pdf"), deflate=True)
    doc.close()


if __name__ == "__main__":
    make_docx()
    make_xlsx()
    make_pdf()
    make_two_column_header_straddle_pdf()
    print("fixtures written to", HERE)
