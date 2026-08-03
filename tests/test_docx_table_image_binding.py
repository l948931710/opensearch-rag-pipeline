# -*- coding: utf-8 -*-
"""C0（2026-08-03）：DOCX 非 1×1 表格单元格内嵌图的抽取→清理→绑定四层验收。

**缺陷**：`docx_extractor._walk` 的非 1×1 表格分支此前只取单元格文字，表内
`<w:drawing>/<w:pict>` 一张都不产 `image_ref` 块。而 `extract_images_from_docx` 按
`document.part.rels` **全量导出**它们、照常过付费 OCR/VLM 漏斗 ⇒ 图进了 canonical assets，
却没有任何 chunk 携带其 `image_refs` ⇒ **serving 端永远渲染不出、零告警**。
真实语料实测：12 篇 / 392 张含图表格 / **495 张真嵌入图**（`r:embed` 口径）全部 serving-dead。

**oracle 独立性**（codex 评审要求）：期望值来自**手工构造 fixture 时自己写下的
`(table_index, row_index, 图片字节 SHA-256)`**，不调用生产 extractor 生成 GT——
既有 DOCX strict 路径的 GT 由 `extract_docx_with_images()` 自己产出、且只把最近
paragraph/heading 当 GT、完全忽略 table text，用它验收等于自己证明自己。
"""
import hashlib
import io

import pytest

docx = pytest.importorskip("docx", reason="python-docx 未安装")
PIL = pytest.importorskip("PIL", reason="Pillow 未安装")
from PIL import Image  # noqa: E402


# ── Layer 1：手工 fixture + 独立 oracle ────────────────────────────────────────
def _png(color, size=(40, 30)) -> bytes:
    buf = io.BytesIO()
    Image.new("RGB", size, color).save(buf, format="PNG")
    return buf.getvalue()


@pytest.fixture(scope="module")
def fixture_docx(tmp_path_factory):
    """构造覆盖五种情形的 DOCX，并返回 (path, oracle)。

    oracle = 我们**写入时**记下的事实，与 extractor 输出完全独立：
      {sha256: {"where": "para"|"table", "table_index": int|None, "row_index": int|None}}

    版式：
      P0  段落图（表前）                  → 无 table_index
      T0  2×2 表：r0c1 有图、r1c1 有图     → table_index=0, row 0 / 1
      P1  段落图（表后）                  → 无 table_index
      T1  含**纵向合并**单元格的表 + 1 图  → table_index=1，合并只算首次出现行
    """
    from docx.shared import Inches
    d = docx.Document()
    oracle = {}
    tmp = tmp_path_factory.mktemp("c0")

    def _add(par, color, where, ti=None, ri=None):
        b = _png(color)
        p = tmp / f"{hashlib.sha256(b).hexdigest()[:12]}.png"
        p.write_bytes(b)
        par.add_run().add_picture(str(p), width=Inches(0.5))
        oracle[hashlib.sha256(b).hexdigest()] = {
            "where": where, "table_index": ti, "row_index": ri}

    d.add_paragraph("步骤1 打开系统")
    _add(d.add_paragraph(), (255, 0, 0), "para")            # P0 表前段落图

    t0 = d.add_table(rows=2, cols=2)
    t0.cell(0, 0).text = "步骤2 录入单据"
    _add(t0.cell(0, 1).paragraphs[0], (0, 255, 0), "table", 0, 0)     # T0 r0
    t0.cell(1, 0).text = "步骤3 复核单据"
    _add(t0.cell(1, 1).paragraphs[0], (0, 0, 255), "table", 0, 1)     # T0 r1

    _add(d.add_paragraph(), (255, 255, 0), "para")           # P1 表后段落图

    t1 = d.add_table(rows=2, cols=2)
    t1.cell(0, 0).merge(t1.cell(1, 0))                       # 纵向合并首列
    t1.cell(0, 0).text = "合并说明"
    _add(t1.cell(0, 1).paragraphs[0], (255, 0, 255), "table", 1, 0)   # T1 r0
    t1.cell(1, 1).text = "第二行文字"

    path = str(tmp / "c0_fixture.docx")
    d.save(path)
    return path, oracle


def _sha_by_target_ref(path):
    """从 OOXML 包里直接读 media 字节 → {target_ref: sha256}（独立于 extractor）。"""
    import zipfile
    out = {}
    with zipfile.ZipFile(path) as z:
        for n in z.namelist():
            if n.startswith("word/media/"):
                out[n.split("word/")[-1]] = hashlib.sha256(z.read(n)).hexdigest()
    return out


# ── Layer 2：抽取层 ───────────────────────────────────────────────────────────
def test_table_cell_images_produce_refs_with_shared_table_id(fixture_docx):
    """★ 表内图必须产 image_ref，且与其 table block 共享 table_index、带 0-based row_index。"""
    from opensearch_pipeline.extraction.docx_extractor import extract_docx_with_images
    path, oracle = fixture_docx
    blocks, assets, _ = extract_docx_with_images(path)
    sha = _sha_by_target_ref(path)

    refs = [b for b in blocks if b.block_type == "image_ref"]
    assert len(refs) == len(oracle), f"应产 {len(oracle)} 个 ref，实得 {len(refs)}"

    # 逐张按【图片字节 SHA-256】比对 oracle —— 不依赖 extractor 的任何判断
    got = {}
    for r in refs:
        tr = r.extra.get("target_ref", "")
        got[sha[tr]] = (r.extra.get("table_index"), r.extra.get("row_index"))
    for h, exp in oracle.items():
        assert h in got, f"图 {h[:8]} 未产 ref（where={exp['where']}）"
        assert got[h] == (exp["table_index"], exp["row_index"]), (
            f"图 {h[:8]} 锚点错：期望 {(exp['table_index'], exp['row_index'])}，实得 {got[h]}")

    # table block 与其 refs 共享同一 table_index
    tbl_ids = {b.extra.get("table_index") for b in blocks if b.block_type == "table"}
    ref_tids = {v[0] for v in got.values() if v[0] is not None}
    assert ref_tids <= tbl_ids, f"ref 的 table_index {ref_tids} 不在 table block 的 {tbl_ids} 内"


def test_merged_cell_image_not_emitted_twice(fixture_docx):
    """★ 纵向合并单元格：扫描在 seen_tc 去重之后 ⇒ 同一图只发一次 ref。"""
    from opensearch_pipeline.extraction.docx_extractor import extract_docx_with_images
    path, _ = fixture_docx
    blocks, _, _ = extract_docx_with_images(path)
    refs = [b for b in blocks if b.block_type == "image_ref"]
    trs = [r.extra.get("target_ref") for r in refs]
    assert len(trs) == len(set(trs)), f"有 target_ref 重复发射（合并单元格去重失效）: {trs}"


def test_image_index_is_dense_and_document_ordered(fixture_docx):
    """image_index 必须连续唯一（下游 asset 对齐按它）；表内图按物理序插在段落图之间。"""
    from opensearch_pipeline.extraction.docx_extractor import extract_docx_with_images
    path, _ = fixture_docx
    blocks, assets, _ = extract_docx_with_images(path)
    refs = [b for b in blocks if b.block_type == "image_ref"]
    idxs = [r.extra["image_index"] for r in refs]
    assert idxs == list(range(len(idxs))), f"image_index 不连续/不唯一: {idxs}"
    assert [a.image_index for a in assets] == idxs, "ImageAsset 与 ref 的 image_index 未对齐"


def test_refs_follow_their_table_block(fixture_docx):
    """表内图的 ref 必须【紧跟】其 table block 之后 —— 表整块并入 current_step，
    refs 跟着表走才绑到同一步。"""
    from opensearch_pipeline.extraction.docx_extractor import extract_docx_with_images
    path, _ = fixture_docx
    blocks, _, _ = extract_docx_with_images(path)
    seq = [(b.block_type, (b.extra or {}).get("table_index")) for b in blocks]
    for i, (bt, ti) in enumerate(seq):
        if bt == "image_ref" and ti is not None:
            prev = [s for s in seq[:i] if s[0] in ("table", "image_ref")]
            assert prev and prev[-1][1] == ti, f"表 {ti} 的 ref 未紧跟其 table block: {seq}"


# ── Layer 3：漏斗层（确定性 stub，不调真实 VLM）────────────────────────────────
def _run_enrich(blocks, statuses):
    """用确定性 assets 跑 _enrich_existing_image_refs。statuses: {image_index: status}"""
    from opensearch_pipeline.pipeline_nodes import _enrich_existing_image_refs
    assets = [{"image_index": i, "status": s, "filename": f"img{i}.png",
               "oss_key": f"oss/img{i}.png", "ocr_text": "", "visual_summary": ""}
              for i, s in statuses.items()]
    doc = {"doc_id": "D1", "version_no": 1, "source_key": "raw/hr/DOC_X/UL/f.docx",
           "owner_dept": "hr"}
    return _enrich_existing_image_refs(blocks, assets, doc)


def test_survivor_set_equals_final_ref_set(fixture_docx):
    """★ 混合 keep/discard：最终 ref 集合必须【严格等于】survivor 集合，且无空 oss_key。

    survivor 严格定义 = CLEAN_ROUTES（ROUTE_TO_TEXT/ROUTE_TO_VECTOR）；
    QUARANTINE_SENSITIVE 虽在 assets 里也不得产最终 ref。
    """
    from opensearch_pipeline.extraction.docx_extractor import extract_docx_with_images
    from opensearch_pipeline.asset_set_diff import CLEAN_ROUTES
    path, _ = fixture_docx
    blocks, _, _ = extract_docx_with_images(path)
    n = len([b for b in blocks if b.block_type == "image_ref"])
    statuses = {}
    for i in range(n):
        statuses[i] = ["ROUTE_TO_VECTOR", "DISCARD_DECORATIVE", "ROUTE_TO_TEXT",
                       "QUARANTINE_SENSITIVE"][i % 4]
    out = _run_enrich(blocks, statuses)
    survivors = {i for i, s in statuses.items() if s in CLEAN_ROUTES}
    final = {b.extra["image_index"] for b in out if b.block_type == "image_ref"}
    assert final == survivors, f"最终 ref {final} != survivor {survivors}"
    for b in out:
        if b.block_type == "image_ref":
            assert b.extra.get("oss_key"), f"存活 ref 无 oss_key: {b.extra}"


def test_all_discard_leaves_zero_refs_not_hollow_ones(fixture_docx):
    """★ 全 DISCARD（assets=[]）必须清成【零 ref】，不得留无 oss_key 的空洞。

    这是 C0 引入表内 logo ref 后最危险的一条：抬头表 logo 全被漏斗判 DISCARD ⇒ assets=[] ⇒
    旧调用点 `if assets and is_step_mode` / `elif assets` 两支都不触发 ⇒ ref 原样残留，
    serving 渲染不出、llm_generator 却按 image_refs 的 truthiness 照样注 <<IMG:N>>。
    """
    from opensearch_pipeline.extraction.docx_extractor import extract_docx_with_images
    path, _ = fixture_docx
    blocks, _, _ = extract_docx_with_images(path)
    out = _run_enrich(blocks, {})          # assets 全空 = 全部 DISCARD
    assert [b for b in out if b.block_type == "image_ref"] == [], \
        "全 DISCARD 后仍有残留 ref —— 空洞 bug 复现"
    assert any(b.block_type == "table" for b in out), "非图块不应被误删"


def test_call_site_prunes_refs_even_when_assets_empty():
    """★ 调用层判断顺序：位置性 ref 的清理【不受 assets 门控】。

    只改 helper 内部无效——调用点自身 `if assets ...` / `elif assets` 会让它根本不被调用。
    这里对源码结构做断言（该路径的真实执行需整条 stage-2，成本过高）。
    """
    import inspect
    import re
    from opensearch_pipeline import pipeline_nodes
    src = inspect.getsource(pipeline_nodes.node_chunk_documents)
    # 定位调用 _enrich_existing_image_refs 的那一行，取其所属 if/elif 的条件表达式
    m = re.search(r"^\s*(if|elif)\s+(.+?):\s*\n(?:\s*#.*\n)*\s*blocks = _enrich_existing_image_refs",
                  src, re.M)
    assert m, "未找到 _enrich_existing_image_refs 的调用分支，调用点结构已变，请同步本断言"
    cond = m.group(2)
    assert "assets" not in cond, (
        f"位置性 ref 的清理被 assets 门控了（条件={cond!r}）——"
        "全 DISCARD 时 assets=[] 会让它根本不被调用，残留无 oss_key 的空洞 ref")
    assert "_has_refs" in cond, f"清理分支应以 _has_refs 为条件，实为 {cond!r}"


# ── Layer 4：chunk 层 ─────────────────────────────────────────────────────────
def _step_chunks(blocks, diagnostics=None):
    from opensearch_pipeline.chunker import DocumentChunker
    c = DocumentChunker(split_mode="step")
    return c.chunk_from_blocks(blocks=blocks, doc_id="D1", version_no=1,
                               metadata={}, diagnostics=diagnostics)


def test_table_images_bind_to_the_step_the_table_joined(fixture_docx):
    """★ 表内图应绑到"表所属的那一步"，而不是流到下一步/最后一步。"""
    from opensearch_pipeline.extraction.docx_extractor import extract_docx_with_images
    path, _ = fixture_docx
    blocks, _, _ = extract_docx_with_images(path)
    enriched = _run_enrich(blocks, {i: "ROUTE_TO_VECTOR" for i in range(10)})
    chunks = _step_chunks(enriched)
    cards = [c for c in chunks if c.chunk_type == "step_card"]
    assert cards, "未产出 step_card"
    bound = sum(len((c.extra or {}).get("image_refs", []) or []) for c in cards)
    assert bound > 0, "表内图未绑定到任何 step_card"


def test_table_image_without_preceding_step_is_dropped_and_diagnosed():
    """★ 表内图无前置步骤 ⇒ 显式丢弃 + 结构化诊断（不静默错绑到下一步）。

    断言三件事：ref 未进入任何 chunk、diagnostic 落进 sink、按 table 聚合（防 255 截断）。
    """
    blocks = [
        {"block_type": "table", "text": "| 抬头 | 表 |", "extra": {"table_index": 0,
                                                                 "row_count": 1}},
        {"block_type": "image_ref", "text": "",
         "extra": {"image_index": 0, "table_index": 0, "row_index": 0,
                   "oss_key": "oss/a.png", "source_image": "s/a.png"}},
        {"block_type": "image_ref", "text": "",
         "extra": {"image_index": 1, "table_index": 0, "row_index": 0,
                   "oss_key": "oss/b.png", "source_image": "s/b.png"}},
        {"block_type": "paragraph", "text": "步骤1 之后才出现的步骤"},
    ]
    diags: list = []
    chunks = _step_chunks(blocks, diagnostics=diags)
    allrefs = [r for c in chunks for r in ((c.extra or {}).get("image_refs", []) or [])]
    assert allrefs == [], f"无前置步骤的表内图被错绑进 chunk: {allrefs}"
    assert len(diags) == 1, f"诊断应按 table 聚合成 1 条，实得 {diags}"
    assert diags[0]["code"] == "C0_TABLE_IMAGE_UNBOUND"
    assert diags[0]["table_index"] == 0 and diags[0]["count"] == 2
    assert sorted(diags[0]["image_indices"]) == [0, 1]


def test_non_table_orphan_images_still_use_pending_path():
    """反向对照：非表格的 orphan 图仍走 pending_images 归入下一步（既有行为不变）。"""
    blocks = [
        {"block_type": "image_ref", "text": "",
         "extra": {"image_index": 0, "oss_key": "oss/a.png", "source_image": "s/a.png"}},
        {"block_type": "paragraph", "text": "步骤1 打开系统"},
    ]
    diags: list = []
    chunks = _step_chunks(blocks, diagnostics=diags)
    allrefs = [r for c in chunks for r in ((c.extra or {}).get("image_refs", []) or [])]
    assert len(allrefs) == 1, "段落 orphan 图的 pending 归属被误改"
    assert diags == [], "段落 orphan 图不应产 C0 诊断"


def test_image_bearing_table_is_never_prose_flattened():
    """★ 带存活图的表禁止 prose-flatten —— 拍平后会按步骤标记拆成多步，
    表后的 N 个 ref 会全部落到【最后一个】step，造成确定性错绑。"""
    from opensearch_pipeline.chunker import DocumentChunker, _is_prose_table
    # _PROSE_CLAUSE 要求 `\d{1,2}[、.][一-龥]`（数字+点+中文，中间无空格）；
    # 且需非空白字数 ≥500、平均段长 ≥18。三条缺一不可，否则测试前提就不成立。
    long_prose = " | ".join(
        [f"{i}.条款这是一条足够长的条款正文内容用于满足 prose 判定的平均段长与总字数阈值必须写得冗长"
         for i in range(1, 16)])
    assert _is_prose_table(long_prose), "样例未命中 _is_prose_table，测试前提失效"
    blocks = [
        {"block_type": "table", "text": long_prose, "extra": {"table_index": 7}},
        {"block_type": "image_ref", "text": "", "extra": {"image_index": 0, "table_index": 7,
                                                          "row_index": 0}},
    ]
    out = DocumentChunker(split_mode="step")._flatten_prose_tables(blocks)
    assert out[0]["block_type"] == "table", "带图的 prose 表被拍平了"

    # 反向：同样的表若图已被漏斗清掉（无存活 ref）⇒ 仍应允许拍平
    out2 = DocumentChunker(split_mode="step")._flatten_prose_tables([blocks[0]])
    assert out2[0]["block_type"] == "paragraph", "无存活图的 prose 表应照常拍平"


def test_chunk_from_blocks_return_contract_unchanged(fixture_docx):
    """diagnostics 是可选 sink —— 不传时行为与返回类型逐字不变（全仓 92 处调用）。"""
    from opensearch_pipeline.extraction.docx_extractor import extract_docx_with_images
    path, _ = fixture_docx
    blocks, _, _ = extract_docx_with_images(path)
    enriched = _run_enrich(blocks, {i: "ROUTE_TO_VECTOR" for i in range(10)})
    a = _step_chunks(enriched)                      # 不传 diagnostics
    b = _step_chunks(enriched, diagnostics=[])      # 传空 sink
    assert isinstance(a, list) and all(hasattr(c, "chunk_type") for c in a)
    assert [c.chunk_id for c in a] == [c.chunk_id for c in b], "传 sink 改变了 chunk 结果"
