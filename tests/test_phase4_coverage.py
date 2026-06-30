# -*- coding: utf-8 -*-
"""tests/test_phase4_coverage.py — Phase-4 dim1/2 fixes.

SF-2: _normalize_image_refs must preserve the xlsx same-anchor contract keys (filename/anchor_row)
      across the RDS→serving roundtrip (CLAUDE.md load-bearing contract).
DC-3: DOCX table extraction must dedup merged cells (gridSpan/vMerge) — python-docx repeats the same
      <w:tc> across spanned grid positions/rows, which otherwise duplicates the merged cell's text.
"""
import os
import tempfile

import pytest


# ── SF-2: serving image_refs roundtrip preserves filename + anchor_row ──

def test_normalize_image_refs_preserves_filename_and_anchor_row():
    from opensearch_pipeline.retriever import _normalize_image_refs
    out = _normalize_image_refs([
        {"oss_key": "k1", "filename": "step_p1_img0003.png", "anchor_row": 7, "visual_summary": "v"},
    ])
    assert out[0]["filename"] == "step_p1_img0003.png"
    assert out[0]["anchor_row"] == 7
    # the pre-existing contract keys are still emitted
    for k in ("oss_key", "source_image", "visual_summary", "ocr_text", "caption", "order", "image_index"):
        assert k in out[0]


def test_normalize_image_refs_safe_defaults_when_absent():
    from opensearch_pipeline.retriever import _normalize_image_refs
    out = _normalize_image_refs([{"oss_key": "k2"}])  # non-xlsx ref, no filename/anchor_row
    assert out[0]["filename"] == ""
    assert out[0]["anchor_row"] is None


def test_normalize_image_refs_accepts_json_string():
    import json
    from opensearch_pipeline.retriever import _normalize_image_refs
    out = _normalize_image_refs(json.dumps([{"oss_key": "k", "filename": "f.png", "anchor_row": 3}]))
    assert out[0]["filename"] == "f.png" and out[0]["anchor_row"] == 3


# ── DC-3: DOCX merged-cell dedup ──

def _save(doc):
    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    doc.save(path)
    return path


@pytest.fixture
def cleanup_files():
    files = []
    yield files
    for p in files:
        if os.path.exists(p):
            os.remove(p)


def _table_text(blocks):
    return "\n".join(b.text for b in blocks if getattr(b, "block_type", "") == "table")


def test_dc3_horizontal_merge_not_duplicated(cleanup_files):
    """A gridSpan (horizontal) merge must emit the merged text once, not per spanned column."""
    import docx
    from opensearch_pipeline.extraction.docx_extractor import extract_docx
    doc = docx.Document()
    t = doc.add_table(rows=2, cols=3)
    merged = t.cell(0, 0).merge(t.cell(0, 1))  # span cols 0-1 in row 0
    merged.text = "MERGED"
    t.cell(0, 2).text = "C"
    t.cell(1, 0).text = "X"
    t.cell(1, 1).text = "Y"
    t.cell(1, 2).text = "Z"
    path = _save(doc)
    cleanup_files.append(path)

    blocks, _ = extract_docx(path)
    text = _table_text(blocks)
    assert "MERGED" in text
    assert "MERGED | MERGED" not in text          # the bug: duplicated merged cell
    assert "MERGED | C" in text                    # correct: deduped
    assert "X | Y | Z" in text                     # non-merged row unaffected


def test_dc3_vertical_merge_not_duplicated(cleanup_files):
    """A vMerge (vertical) merge must emit the merged text once (top row), not in continuation rows."""
    import docx
    from opensearch_pipeline.extraction.docx_extractor import extract_docx
    doc = docx.Document()
    t = doc.add_table(rows=2, cols=2)
    v = t.cell(0, 0).merge(t.cell(1, 0))  # span rows 0-1 in col 0
    v.text = "VMERGE"
    t.cell(0, 1).text = "P"
    t.cell(1, 1).text = "Q"
    path = _save(doc)
    cleanup_files.append(path)

    blocks, _ = extract_docx(path)
    text = _table_text(blocks)
    assert text.count("VMERGE") == 1               # appears once, not in both rows
    assert "VMERGE | P" in text


def test_dc3_non_merged_table_unaffected(cleanup_files):
    """A plain table without merges keeps its full cell grid (no over-dedup)."""
    import docx
    from opensearch_pipeline.extraction.docx_extractor import extract_docx
    doc = docx.Document()
    t = doc.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "A"
    t.cell(0, 1).text = "B"
    t.cell(1, 0).text = "C"
    t.cell(1, 1).text = "D"
    path = _save(doc)
    cleanup_files.append(path)

    blocks, _ = extract_docx(path)
    text = _table_text(blocks)
    assert "A | B" in text and "C | D" in text


def test_dc3_dedup_in_with_images_path(cleanup_files):
    """The second table loop (extract_docx_with_images) must dedup merged cells too."""
    import docx
    from opensearch_pipeline.extraction.docx_extractor import extract_docx_with_images
    doc = docx.Document()
    t = doc.add_table(rows=1, cols=3)
    merged = t.cell(0, 0).merge(t.cell(0, 1))
    merged.text = "HDR"
    t.cell(0, 2).text = "Z"
    path = _save(doc)
    cleanup_files.append(path)

    blocks, _, _ = extract_docx_with_images(path)
    text = _table_text(blocks)
    assert "HDR | HDR" not in text
    assert "HDR | Z" in text
